"""DOCX -> HTML converter."""

from __future__ import annotations

import hashlib
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import List, Tuple

from docx import Document  # type: ignore
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
from docx.oxml.ns import qn  # type: ignore

from .artifacts import AssetUpload, ConvertedDocument, CssBundle
from .styles import StyleRegistry

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CDN_BASE = "https://cdn.litternetworks.org"
IMAGES_PREFIX = "docs/images"


@dataclass
class DocxConverter:
    registry: StyleRegistry

    def convert(self, source: Path, relative_path: Path) -> ConvertedDocument:
        document = Document(source)
        body_fragments: List[str] = []
        title, subtitle = "", ""
        self._asset_uploads: List[AssetUpload] = []
        self._asset_keys: set[str] = set()

        paragraphs = list(document.paragraphs)
        index = 0
        while index < len(paragraphs):
            paragraph = paragraphs[index]
            para_style = (paragraph.style.name or "").lower() if paragraph.style else ""
            if para_style.startswith("heading"):
                level = self._heading_level(para_style)
                heading_html = self._convert_heading(paragraph.text, level)
                if heading_html:
                    body_fragments.append(heading_html)
                if not title:
                    title = paragraph.text.strip()
                index += 1
                continue

            if para_style == "title" and not title:
                title = paragraph.text.strip()
                index += 1
                continue

            if para_style == "subtitle" and not subtitle:
                subtitle = paragraph.text.strip()
                index += 1
                continue

            list_meta = self._list_metadata(paragraph)
            if list_meta:
                list_html, consumed = self._build_list(paragraphs, index, list_meta)
                body_fragments.append(list_html)
                index += consumed
                continue

            paragraph_html = self._convert_paragraph(paragraph)
            if paragraph_html is not None:
                body_fragments.append(paragraph_html)
            index += 1

        html_body = "\n".join(fragment for fragment in body_fragments if fragment)
        shell_class = self.registry.shell_class()
        body_html = f'<div class="{shell_class}">\n{html_body}\n</div>'
        css_bundle = CssBundle(name=self._css_name(relative_path), rules=self.registry.rules())
        css_href = (
            f"https://cdn.litternetworks.org/docs/styles/{css_bundle.name}.css"
        )
        head_fragment = f'<link rel="stylesheet" href="{css_href}"/>'
        document_html = "\n".join([head_fragment, body_html])
        return ConvertedDocument(
            source=source,
            relative_path=relative_path,
            html=document_html,
            title=title,
            subtitle=subtitle,
            assets=list(self._asset_uploads),
            css_bundle=css_bundle,
        )

    def _css_name(self, relative_path: Path) -> str:
        tokens = list(relative_path.with_suffix("").parts)
        return "-".join(tokens)

    def _convert_heading(self, text: str, level: int) -> str:
        stripped = text.strip()
        if not stripped:
            return ""
        safe_level = min(max(level, 1), 6)
        content = escape(stripped)
        return f"<h{safe_level}>{content}</h{safe_level}>"

    def _convert_paragraph(self, paragraph) -> str | None:  # type: ignore[override]
        content = self._paragraph_content(paragraph)
        alignment = self._paragraph_alignment(paragraph.alignment)
        klass = self.registry.alignment_class(alignment)
        if not content:
            return f'<p class="{klass}"></p>'
        return f'<p class="{klass}">{content}</p>'

    def _paragraph_content(self, paragraph) -> str:
        fragments: List[str] = []
        for child in paragraph._p:
            tag = child.tag
            if tag == qn("w:hyperlink"):
                fragments.append(self._convert_hyperlink(paragraph, child))
            elif tag == qn("w:r"):
                fragments.append(self._convert_run_element(paragraph, child))
            elif tag == qn("w:br"):
                fragments.append("<br/>")
        return "".join(fragments).strip()

    def _convert_hyperlink(self, paragraph, hyperlink) -> str:  # type: ignore[override]
        rel_id = hyperlink.get(qn("r:id"))
        anchor = hyperlink.get(qn("w:anchor"))
        href = ""
        if rel_id and rel_id in paragraph.part.rels:
            href = paragraph.part.rels[rel_id].target_ref
        elif anchor:
            href = f"#{anchor}"
        text = "".join(
            self._convert_run_element(paragraph, run)
            for run in hyperlink.findall(f".//{{{W_NS}}}r")
        )
        safe_href = escape(href, quote=True)
        return f'<a href="{safe_href}">{text}</a>'

    def _convert_run_element(self, paragraph, run_element) -> str:  # type: ignore[override]
        parts: List[str] = []
        for child in run_element:
            if child.tag == qn("w:t"):
                parts.append(escape(child.text or ""))
            elif child.tag == qn("w:tab"):
                parts.append("&emsp;")
            elif child.tag == qn("w:br"):
                parts.append("<br/>")
            elif child.tag == qn("w:drawing"):
                parts.append(self._convert_drawing(paragraph, child))
        run_text = "".join(parts)
        if not run_text:
            return ""

        r_pr = run_element.find("w:rPr", namespaces={"w": W_NS})
        open_tags: List[str] = []
        close_tags: List[str] = []

        def wrap(tag: str) -> None:
            open_tags.append(f"<{tag}>")
            close_tags.insert(0, f"</{tag}>")

        if r_pr is not None:
            if r_pr.find("w:b", namespaces={"w": W_NS}) is not None:
                wrap("strong")
            if r_pr.find("w:i", namespaces={"w": W_NS}) is not None:
                wrap("em")
            if r_pr.find("w:u", namespaces={"w": W_NS}) is not None:
                wrap("u")
            if r_pr.find("w:strike", namespaces={"w": W_NS}) is not None:
                wrap("s")

        return "".join(open_tags) + run_text + "".join(close_tags)

    def _paragraph_alignment(self, alignment) -> str:  # type: ignore[override]
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            return "center"
        if alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            return "right"
        if alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
            return "left"
        return "justify"

    def _heading_level(self, style_name: str) -> int:
        try:
            return int(style_name.replace("heading", ""))
        except ValueError:
            return 1

    def _convert_drawing(self, paragraph, drawing) -> str:
        hlink = next((el for el in drawing.iter() if el.tag == f"{{{A_NS}}}hlinkClick"), None)
        if hlink is not None:
            rel_id = hlink.get(qn("r:id"))
            if rel_id and rel_id in paragraph.part.rels:
                url = paragraph.part.rels[rel_id].target_ref
                safe_url = escape(url, quote=True)
                return (
                    '<div class="video-container">'
                    f'<iframe src="{safe_url}" frameborder="0" allowfullscreen></iframe>'
                    "</div>"
                )
        blip = next((el for el in drawing.iter() if el.tag == f"{{{A_NS}}}blip"), None)
        if blip is None:
            return ""
        rel_id = blip.get(qn("r:embed"))
        if not rel_id or rel_id not in paragraph.part.related_parts:
            return ""
        image_part = paragraph.part.related_parts[rel_id]
        blob = image_part.blob
        digest = hashlib.sha256(blob).hexdigest()
        ext = image_part.partname.ext or ""
        if ext and not ext.startswith("."):
            ext = f".{ext}"
        key = f"{IMAGES_PREFIX}/{digest}{ext}"
        self._register_asset(key, blob, image_part.content_type)
        width_pt, height_pt = self._drawing_dimensions(drawing)
        style_segments = []
        if width_pt:
            style_segments.append(f"width:{width_pt:.2f}pt;")
        if height_pt:
            style_segments.append(f"height:{height_pt:.2f}pt;")
        style_attr = "".join(style_segments)
        src = f"{CDN_BASE}/{key}"
        style_html = f' style="{style_attr}"' if style_attr else ""
        return f'<img src="{src}"{style_html} alt=""/>'

    def _drawing_dimensions(self, drawing) -> Tuple[float | None, float | None]:
        extent = next((el for el in drawing.iter() if el.tag == f"{{{WP_NS}}}extent"), None)
        if extent is None:
            return None, None
        cx = extent.get("cx")
        cy = extent.get("cy")
        width_pt = float(cx) / 12700 if cx else None
        height_pt = float(cy) / 12700 if cy else None
        return width_pt, height_pt

    def _register_asset(self, key: str, blob: bytes, content_type: str) -> None:
        if key in self._asset_keys:
            return
        self._asset_keys.add(key)
        self._asset_uploads.append(AssetUpload(key=key, body=blob, content_type=content_type))

    def _list_metadata(self, paragraph):
        props = paragraph._p.pPr
        if props is None or props.numPr is None or props.numPr.numId is None:
            return None
        num_id = props.numPr.numId.val
        level = props.numPr.ilvl.val if props.numPr.ilvl is not None else 0
        return {"num_id": num_id, "level": level, "type": "ul"}

    def _build_list(self, paragraphs, start_index: int, metadata):
        items: List[str] = []
        index = start_index
        while index < len(paragraphs):
            paragraph = paragraphs[index]
            current_meta = self._list_metadata(paragraph)
            if (
                current_meta is None
                or current_meta["num_id"] != metadata["num_id"]
                or current_meta["level"] != metadata["level"]
            ):
                break
            content = self._paragraph_content(paragraph)
            items.append(f"<li>{content}</li>")
            index += 1

        list_tag = metadata["type"]
        html = [f"<{list_tag}>", *items, f"</{list_tag}>"]
        consumed = index - start_index
        return "\n".join(html), consumed


__all__ = ["DocxConverter"]
